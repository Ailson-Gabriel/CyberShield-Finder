import os
import re
from docx import Document
from valida_cpf import valida_cpf

def encontrar_nomes(caminho_diretorio, nomes):
    """
    Procura pelos nomes nos arquivos .docx em um diretório especificado.

    Argumentos:
        caminho_diretorio (str): O caminho para o diretório contendo o arquivo .docx.
        nomes (lista): Uma lista de nomes a serem procurados no arquivo .docx.

    Retornos:
        lista: Uma lista dos nomes encontrados nos arquivos .docx.
    """
    nomes_encontrados = [] # Inicializa uma lista para armazenar os nomes encontrados
    document = Document(caminho_diretorio) # Abre o arquivo .docx usando a classe Document do módulo docx
    for paragraph in document.paragraphs: # Repete para cada parágrafo no documento .docx
        for nome in nomes: # Repete para cada nome na lista de nomes fornecida
            if nome.lower() in paragraph.text.lower(): # Verifica se o nome (em minúsculo) está presente no texto do parágrafo (em minúsculo)
                nomes_encontrados.append(nome) # Adiciona à lista de nomes encontrados
                #print(f"Nome encontrado: {nome} - Arquivo: {arquivo}") # Imprime mensagem indicando o nome encontrado e o arquivo onde foi encontrado
    return nomes_encontrados


def encontrar_cpf(caminho_diretorio):
    """
    Procura por cadeias de 11 números (supostos CPFs) em arquivos .txt em um diretório especificado.
    Valida cada CPF encontrado usando a função valida_cpf de outro arquivo .py.

    Argmuentos:
        caminho_arquivo (str): O caminho para o diretório contendo os arquivos .txt.

    Retorna:
        lista: Uma lista dos CPFs válidos encontrados nos arquivos .docx.
    """
    document = Document(caminho_diretorio) # Abre o arquivo .docx e extrai o texto
    texto_docx = ""
    for paragraph in document.paragraphs: # Repete sobre os parágrafos do documento e concatena seus textos
        texto_docx += paragraph.text + "\n"

    cpfs_potenciais = re.findall(r'\b\d{11}\b', texto_docx) # Encontra todos os conjuntos de 11 números usando expressão regular

    # Valida cada CPF potencial
    cpfs_validos = []
    for cpf in cpfs_potenciais:
        if valida_cpf(cpf): # Chama a função valida_cpf para validar o CPF
            cpfs_validos.append(cpf) # Adiciona o CPF válido à lista de CPFs válidos
            print(f"CPF válido encontrado: {cpf} - Arquivo: {os.path.basename(caminho_diretorio)}") #REMOVER ESTA LINHA APÓS VERSÃO DE TESTES
    return cpfs_validos